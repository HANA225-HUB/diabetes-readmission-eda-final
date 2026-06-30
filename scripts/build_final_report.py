from __future__ import annotations

import hashlib
import json
from pathlib import Path
from textwrap import dedent

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/huahaowen/Documents/Codex/2026-06-24/6-30-23-59-1-pdf")
WORK = ROOT / "work" / "eda_final"
FIG = WORK / "figures"
TAB = WORK / "tables"
RES = WORK / "results"
OUT = ROOT / "outputs"
OUT.mkdir(parents=True, exist_ok=True)

REPORT = OUT / "2462404009_花浩文_探索性数据分析_期末大作业.docx"
CODE_TXT = OUT / "2462404009_花浩文_探索性数据分析_期末大作业_完整代码.txt"
SOURCE_CODE = WORK / "diabetes_readmission_project.py"

FONT_CN = "Songti SC"
FONT_HEAD = "Heiti SC"
FONT_CODE = "Menlo"
INK = "263238"
NAVY = "22577A"
TEAL = "2A9D8F"
GOLD = "E9C46A"
CORAL = "D2644A"
MUTED = "64748B"
LIGHT = "F3F6F8"
MID = "D9E2E8"


audit = json.loads((RES / "data_audit.json").read_text(encoding="utf-8"))
split = json.loads((RES / "split_summary.json").read_text(encoding="utf-8"))
summary = json.loads((RES / "project_summary.json").read_text(encoding="utf-8"))
models = pd.read_csv(TAB / "05_model_comparison.csv")
missing = pd.read_csv(TAB / "01_missingness.csv", index_col=0)
age_stats = pd.read_csv(TAB / "02_age_readmission.csv")
prior_stats = pd.read_csv(TAB / "03_prior_inpatient_readmission.csv")
importance = pd.read_csv(TAB / "06_permutation_importance.csv")
errors = pd.read_csv(TAB / "07_error_analysis.csv")
subgroups = pd.read_csv(TAB / "08_subgroup_performance.csv")
group_cv = pd.read_csv(TAB / "09_group_cv.csv")
capacity = pd.read_csv(TAB / "10_capacity_analysis.csv")
ablation = pd.read_csv(TAB / "11_ablation.csv")
leakage = pd.read_csv(TAB / "12_split_leakage.csv")


def pct(value: float, digits: int = 1) -> str:
    return f"{value:.{digits}%}"


def dec(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}"


def set_run_font(run, name=FONT_CN, size=10.5, bold=None, color=INK, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rpr.rFonts.set(qn(key), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill=LIGHT):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_cm):
    table.autofit = False
    total_dxa = sum(int(width / 2.54 * 1440) for width in widths_cm)
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total_dxa))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), "120")
    tblind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width / 2.54 * 1440)))
        grid.append(col)

    for row in table.rows:
        for index, width in enumerate(widths_cm):
            cell = row.cells[index]
            cell.width = Cm(width)
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(int(width / 2.54 * 1440)))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    trpr.append(node)


def page_number_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run_font(run, FONT_HEAD, 8.5, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.33

    specs = {
        "Title": (28, NAVY, 0, 8),
        "Subtitle": (13, MUTED, 0, 8),
        "Heading 1": (16, NAVY, 16, 8),
        "Heading 2": (13, NAVY, 12, 6),
        "Heading 3": (11.5, "3B5568", 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = FONT_HEAD
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name.startswith("Heading")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.08

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2


def configure_section(section):
    # narrative_proposal preset with named academic_A4 override.
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.65)
    section.right_margin = Cm(2.45)
    section.header_distance = Cm(1.05)
    section.footer_distance = Cm(1.05)
    section.different_first_page_header_footer = True


def set_running_furniture(section):
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("探索性数据分析 · 期末大作业")
    set_run_font(r, FONT_HEAD, 8.5, color=MUTED)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("—  ")
    set_run_font(r, FONT_HEAD, 8.5, color=MUTED)
    page_number_field(fp)
    r = fp.add_run("  —")
    set_run_font(r, FONT_HEAD, 8.5, color=MUTED)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_list(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.left_indent = Cm(0.95)
        p.paragraph_format.first_line_indent = Cm(-0.48)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(item)
        set_run_font(r)


def add_note(doc, title, text, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.right_indent = Cm(0.35)
    shade_paragraph(p)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), color)
    pbdr.append(left)
    ppr.append(pbdr)
    r = p.add_run(f"{title}  ")
    set_run_font(r, FONT_HEAD, 10.4, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r, size=10)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED)


def add_figure(doc, filename, width_cm, caption):
    path = FIG / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(path), width=Cm(width_cm))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption.split("  ", 1)[-1])
    add_caption(doc, caption)


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.28)
    p.paragraph_format.right_indent = Cm(0.18)
    p.paragraph_format.line_spacing = 1.0
    shade_paragraph(p, "F4F4F4")
    r = p.add_run(dedent(code).strip())
    set_run_font(r, FONT_CODE, 7.7, color="263238")


def add_table(doc, headers, rows, widths_cm, alignments=None, font_size=9.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths_cm)
    header = table.rows[0]
    repeat_header(header)
    for index, value in enumerate(headers):
        set_cell_shading(header.cells[index], NAVY)
        p = header.cells[index].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(value))
        set_run_font(r, FONT_HEAD, 9.1, bold=True, color="FFFFFF")
    for row_no, values in enumerate(rows):
        cells = table.add_row().cells
        if row_no % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "F8FAFB")
        for index, value in enumerate(values):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            if alignments:
                p.alignment = alignments[index]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
        set_table_geometry(table, widths_cm)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table


def add_stat_strip(doc, stats):
    table = doc.add_table(rows=1, cols=len(stats))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    repeat_header(table.rows[0])
    widths = [15.9 / len(stats)] * len(stats)
    set_table_geometry(table, widths)
    for index, (value, label) in enumerate(stats):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "EEF5F5" if index % 2 == 0 else "F8F3E5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(value)
        set_run_font(r, FONT_HEAD, 15, bold=True, color=NAVY)
        p.add_run("\n")
        r = p.add_run(label)
        set_run_font(r, FONT_HEAD, 8.5, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def build_workflow_figure():
    path = FIG / "00_research_workflow.png"
    font_path = "/System/Library/Fonts/STHeiti Medium.ttc"
    light_path = "/System/Library/Fonts/STHeiti Light.ttc"
    title_font = ImageFont.truetype(font_path, 29)
    body_font = ImageFont.truetype(light_path, 20)
    small_font = ImageFont.truetype(light_path, 16)
    canvas = Image.new("RGB", (1680, 380), "white")
    draw = ImageDraw.Draw(canvas)
    colors = ["#E8EEF2", "#DFF0EC", "#FFF4D6", "#E5EDF5", "#F5E6E1"]
    titles = ["原始数据", "质量控制", "三层特征", "双模型实验", "临床评价"]
    details = [
        "101,766 条住院记录\n50 个字段",
        "缺失审计 · 排除死亡/临终\n患者级拆分防泄漏",
        "Raw · Cleaned\nEngineered",
        "Logistic Regression\nHistGradientBoosting",
        "PR-AUC · F2 · 校准\n解释 · 误差 · 亚组",
    ]
    x_positions = [35, 365, 695, 1025, 1355]
    for index, x in enumerate(x_positions):
        draw.rounded_rectangle((x, 85, x + 285, 290), radius=18, fill=colors[index], outline="#CBD5DC", width=2)
        box = draw.textbbox((0, 0), titles[index], font=title_font)
        draw.text((x + (285 - (box[2] - box[0])) / 2, 112), titles[index], fill="#22577A", font=title_font)
        lines = details[index].split("\n")
        for line_no, line in enumerate(lines):
            fnt = body_font if line_no == 0 else small_font
            box = draw.textbbox((0, 0), line, font=fnt)
            draw.text((x + (285 - (box[2] - box[0])) / 2, 188 + line_no * 37), line, fill="#3B5568", font=fnt)
        if index < 4:
            draw.line((x + 290, 188, x + 326, 188), fill="#64748B", width=5)
            draw.polygon([(x + 326, 188), (x + 309, 178), (x + 309, 198)], fill="#64748B")
    canvas.save(path)


def build_report():
    build_workflow_figure()
    doc = Document()
    doc.core_properties.author = "花浩文"
    doc.core_properties.title = "糖尿病患者30天再入院预测中的特征工程研究"
    doc.core_properties.subject = "探索性数据分析期末大作业"
    configure_styles(doc)
    configure_section(doc.sections[0])
    set_running_furniture(doc.sections[0])

    # Editorial-cover pattern, deliberately simple for a formal academic report.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    r = p.add_run("探索性数据分析")
    set_run_font(r, FONT_HEAD, 15, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("特征工程如何改变")
    set_run_font(r, FONT_HEAD, 28, bold=True, color=NAVY)
    p.add_run("\n")
    r = p.add_run("糖尿病患者 30 天再入院预测")
    set_run_font(r, FONT_HEAD, 28, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(32)
    r = p.add_run("基于 130 家美国医院住院记录的患者级对照实验")
    set_run_font(r, FONT_HEAD, 13.5, color="3B5568")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("期末大作业研究报告")
    set_run_font(r, FONT_HEAD, 13, bold=True)

    for label, value in [
        ("学生姓名", "花浩文"),
        ("学号", "2462404009"),
        ("课程名称", "探索性数据分析"),
        ("完成时间", "2026 年 6 月"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        r = p.add_run(f"{label}：")
        set_run_font(r, FONT_HEAD, 11, bold=True, color=MUTED)
        r = p.add_run(value)
        set_run_font(r, size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    r = p.add_run("数据来源：Diabetes 130-US Hospitals for Years 1999-2008")
    set_run_font(r, size=9.3, color=MUTED)
    doc.add_page_break()

    doc.add_heading("摘  要", level=1)
    abstract = (
        "30 天内非计划再入院是评价医疗质量与资源配置的重要指标。本文使用美国 130 家医院及综合医疗网络在"
        "1999-2008 年间收集的糖尿病住院记录，研究特征清洗、语义编码与领域衍生变量如何影响机器学习的再入院"
        "预测性能。原始数据含 101,766 条记录和 50 个字段；排除死亡、临终关怀及无效性别记录后，保留 99,340 条"
        "住院记录、69,987 名患者，30 天再入院率为 11.39%。为防止同一患者的多次住院同时进入训练集和测试集，"
        "本文采用患者级训练/验证/测试划分，并在完全相同的划分上比较 Raw、Cleaned、Engineered 三套特征方案与"
        "逻辑回归、直方图梯度提升两类模型。由于正类稀少，PR-AUC 被设为排序主指标；面向筛查应用时，使用验证集"
        "选择最大化 F2 的阈值。结果显示，Raw + HistGradientBoosting 的 PR-AUC 最高（0.210），但经语义清洗后的"
        "Cleaned + HistGradientBoosting 在验证集 F2 上最优，并在独立测试集达到 ROC-AUC 0.664、PR-AUC 0.208、"
        "召回率 81.54% 与 F2 0.418。与 Raw 方案相比，清洗方案提高召回率约 7.9 个百分点，却降低精确率并增加假阳性；"
        "完整 Engineered 方案没有继续带来稳定增益。置换重要性表明，既往住院次数、出院去向、支付信息可得性、"
        "胰岛素使用和既往急诊次数最具预测贡献。研究说明：高质量特征工程的价值不只是增加变量，而在于明确预测时点、"
        "处理缺失语义、控制患者级泄漏，并根据实际筛查目标选择合适的性能权衡。"
    )
    add_body(doc, abstract)
    p = doc.add_paragraph()
    r = p.add_run("关键词：")
    set_run_font(r, FONT_HEAD, 10.5, bold=True, color=NAVY)
    r = p.add_run("探索性数据分析；特征工程；30 天再入院；类别不平衡；梯度提升；患者级数据划分")
    set_run_font(r)

    doc.add_heading("核心结论", level=1)
    add_stat_strip(doc, [
        (f"{audit['analysis_rows']:,}", "分析记录"),
        (pct(audit["positive_rate"]), "30 天再入院率"),
        (dec(summary["best_metrics"]["roc_auc"]), "ROC-AUC"),
        (pct(summary["best_metrics"]["recall"]), "筛查召回率"),
    ])
    add_list(doc, [
        "患者级拆分是本研究最重要的防泄漏设计：23.35% 的患者有多次住院，最多 40 次；若按行随机拆分，模型可能记住患者历史。",
        "语义清洗提升了筛查灵敏度而非整体排序能力。Cleaned 方案把召回率从 73.62% 提高到 81.54%，但精确率从 15.18% 降至 14.20%。",
        "更复杂的特征不一定更好。加入诊断大类、强度比率和药物聚合后的 Engineered 方案未显著提高 PR-AUC，提示冗余与噪声会抵消理论收益。",
        "既往住院次数远高于其他特征的重要性；无既往住院者再入院率为 8.59%，过去一年住院 5 次及以上者达到 37.13%。",
        "模型适合出院时风险筛查和资源优先级排序，不应被解释为因果模型，也不能直接替代临床判断。",
    ])

    doc.add_heading("目录", level=1)
    for item in [
        "1  研究背景与研究问题", "2  数据来源、对象与预测目标", "3  数据质量审计与预处理",
        "4  探索性数据分析", "5  特征工程方案", "6  模型与实验设计", "7  实验结果与特征工程影响",
        "8  模型解释、误差与亚组审计", "9  讨论与局限", "10  结论与建议", "参考文献", "附录",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, FONT_HEAD, 10.5, color="3B5568")

    doc.add_page_break()
    doc.add_heading("1  研究背景与研究问题", level=1)
    doc.add_heading("1.1  研究背景", level=2)
    add_body(doc, "糖尿病患者常伴随心血管、肾脏和感染等多种共病，再入院既可能反映病情复杂，也可能提示出院计划、药物管理和随访衔接不足。30 天内再入院是一个低发生率但高成本的结局：若模型能够在出院时识别高风险患者，医院可以优先安排电话随访、药师复核、门诊预约和健康教育。")
    add_body(doc, "然而，本数据集包含大量类别编码、极高缺失字段、重复患者、几十种药物状态与不平衡目标。直接把原始表格交给分类器会产生三类问题：模型可能把无意义的编号当连续数值；缺失值可能既代表“未检测”又代表“未记录”；按住院记录随机拆分会让同一患者出现在训练集和测试集，形成隐蔽的信息泄漏。因此，本作业的核心不只是比较分类算法，而是系统检验特征工程如何改变模型的排序、筛查与解释表现。")

    doc.add_heading("1.2  研究问题", level=2)
    add_list(doc, [
        "在患者级无泄漏划分下，逻辑回归与非线性梯度提升对 30 天再入院的预测能力如何？",
        "原始特征、语义清洗特征与领域工程特征分别对 ROC-AUC、PR-AUC、召回率、精确率和 F2 产生什么影响？",
        "哪些住院历史、护理路径、实验室检查和用药变量对模型贡献最大？",
        "验证集阈值优化带来怎样的假阴性—假阳性权衡？模型在不同性别、种族和年龄组上的表现是否一致？",
    ], numbered=True)

    doc.add_heading("1.3  研究流程", level=2)
    add_figure(doc, "00_research_workflow.png", 15.6, "图 1  本研究从原始数据到临床评价的完整流程")
    add_note(doc, "分析定位", "本文预测时点设为出院时，因此出院去向可以作为特征；若应用目标改为入院早期预警，必须删除出院去向、最终住院天数等只有住院结束后才能获得的变量。")

    doc.add_heading("2  数据来源、对象与预测目标", level=1)
    doc.add_heading("2.1  数据来源", level=2)
    add_body(doc, "数据集为 Diabetes 130-US Hospitals for Years 1999-2008，来自美国 130 家医院及综合医疗服务网络。每行是一条糖尿病患者住院记录，包含人口学特征、入院与出院路径、住院过程、诊断、实验室检测、糖尿病药物变化及再入院结局。IDS_mapping.csv 提供 admission_type_id、discharge_disposition_id 和 admission_source_id 的文字含义。")
    add_table(doc, ["维度", "数据口径", "建模作用"], [
        ("记录单位", "一次住院 encounter", "预测该次出院后的 30 天再入院"),
        ("时间范围", "1999-2008 年", "跨机构历史临床管理数据"),
        ("原始规模", "101,766 行 × 50 列", "约 10 万条住院记录"),
        ("患者标识", "patient_nbr", "仅用于患者级划分，不进入模型"),
        ("目标变量", "readmitted == '<30'", "二分类正类；>30 与 NO 合并为负类"),
        ("预测时点", "出院时", "允许使用住院过程和出院去向"),
    ], [2.5, 5.0, 8.4])
    add_caption(doc, "表 1  数据对象与预测任务定义")

    doc.add_heading("2.2  样本筛选", level=2)
    add_body(doc, f"原始数据共 {audit['original_rows']:,} 条记录。死亡或进入临终关怀的患者没有通常意义上的再入院机会，若保留会把结局不可发生误当成低风险，因此排除 discharge_disposition_id 为 11、13、14、19、20、21 的 {audit['excluded_death_hospice']:,} 条记录；另排除性别为 Unknown/Invalid 的 {audit['excluded_invalid_gender']} 条记录。最终分析集为 {audit['analysis_rows']:,} 条记录、{audit['unique_patients']:,} 名患者。")
    add_figure(doc, "01_target_distribution.png", 13.3, "图 2  分析样本中的目标变量分布")
    add_body(doc, f"正类共有 {audit['positive_count']:,} 条，占 {pct(audit['positive_rate'], 2)}。由于随机分类器的 PR-AUC 基准约等于正类率，单独报告准确率会严重高估性能：全部预测为“未再入院”即可获得 88.61% 的准确率，却无法识别任何高风险患者。")

    doc.add_heading("2.3  数据伦理与隐私边界", level=2)
    add_body(doc, "数据已使用匿名患者编号，不包含姓名和直接联系方式。本文只进行群体层面的模型研究，不尝试重识别患者。种族、性别和年龄既可能反映健康差异，也可能成为代理变量，因此只进行描述性亚组审计，不把组间差异解释为生物学因果，更不建议将模型输出用于拒绝服务或降低资源配置。")

    doc.add_heading("3  数据质量审计与预处理", level=1)
    doc.add_heading("3.1  缺失值结构", level=2)
    add_figure(doc, "02_missingness.png", 13.5, "图 3  原始数据中缺失率最高的 12 个字段")
    missing_rows = []
    strategy = {
        "weight": "缺失 96.86%，不进入模型",
        "max_glu_serum": "将缺失解释为未检测，并构造 tested/abnormal 指示",
        "A1Cresult": "将缺失解释为未检测，并构造 tested/abnormal 指示",
        "medical_specialty": "Missing 单独成类，低频专科合并为 Other",
        "payer_code": "不直接使用编码，只保留是否记录 payer_available",
        "race": "缺失单独成类，不用众数覆盖",
        "diag_3": "缺失单独成类，再映射为诊断大类",
        "diag_2": "缺失单独成类，再映射为诊断大类",
    }
    for field in strategy:
        missing_rows.append((field, pct(float(missing.loc[field, "missing_rate"]), 2), strategy[field]))
    add_table(doc, ["字段", "缺失率", "处理策略与理由"], missing_rows, [3.7, 2.2, 10.0])
    add_caption(doc, "表 2  高缺失字段的处理策略")
    add_body(doc, "缺失并不总是随机噪声。HbA1c 和最高血糖的空值大多意味着本次住院未做该项检测；支付方和专科是否被记录也可能与医院流程相关。因而本文尽量保留“未检测/未记录”的信息，而不是简单使用众数或均值填充。")

    doc.add_heading("3.2  编码、异常与泄漏控制", level=2)
    add_list(doc, [
        "将 admission_type_id、admission_source_id、discharge_disposition_id 作为类别变量，并依据映射表合并为 Emergency、Referral、Facility/Transfer 等临床语义组，避免把编号大小误当作连续关系。",
        "年龄区间转换为中点 age_mid，同时在 Raw 方案中保留原始区间，比较顺序信息是否有益。",
        "ICD-9 诊断码映射为 Diabetes、Circulatory、Respiratory、Digestive、Genitourinary、Neoplasms 等大类，减少数百个稀疏编码。",
        "encounter_id 是记录标识；patient_nbr 只用于划分患者，不作为预测变量。目标 readmitted 及其任何直接派生变量均不进入特征。",
        "所有缺失填补、标准化与类别编码均封装在 sklearn Pipeline 中，只在训练集拟合，验证集和测试集仅执行 transform。",
    ])

    doc.add_heading("3.3  患者级训练、验证与测试划分", level=2)
    add_figure(doc, "08_repeat_patients.png", 12.8, "图 4  患者重复住院记录分布")
    add_body(doc, f"共有 {pct(audit['repeat_patient_rate'], 2)} 的患者出现两次及以上，单个患者最多 {audit['max_encounters_per_patient']} 次。本文先按患者是否曾发生 30 天再入院进行分层，再把患者分为训练、验证与测试集合；患者在三者之间互斥。")
    add_table(doc, ["集合", "患者数", "记录数", "正类率", "用途"], [
        ("训练集", f"{split['train_patients']:,}", f"{split['train_rows']:,}", pct(split["train_rate"], 2), "拟合预处理器与模型"),
        ("验证集", f"{split['validation_patients']:,}", f"{split['validation_rows']:,}", pct(split["validation_rate"], 2), "选择 F2 阈值和最终方案"),
        ("测试集", f"{split['test_patients']:,}", f"{split['test_rows']:,}", pct(split["test_rate"], 2), "一次性报告最终泛化性能"),
    ], [2.5, 2.4, 2.4, 2.4, 6.2])
    add_caption(doc, "表 3  患者级无泄漏数据划分")

    doc.add_heading("4  探索性数据分析", level=1)
    doc.add_heading("4.1  年龄与再入院", level=2)
    add_figure(doc, "03_readmission_by_age.png", 14.2, "图 5  不同年龄区间的 30 天再入院率及 95% 置信区间")
    add_body(doc, "总体上，中老年组的风险高于儿童和青少年，但关系并非严格单调。[20-30) 组的再入院率为 14.31%，可能与样本量较小及疾病亚型差异有关；[70-80) 与 [80-90) 组分别为 12.06% 和 12.57%。年龄适合保留非线性表达，不能只假设每增加一岁风险等幅变化。")

    doc.add_heading("4.2  既往医疗利用", level=2)
    add_figure(doc, "04_readmission_by_prior_inpatient.png", 13.3, "图 6  过去一年住院次数与 30 天再入院率")
    add_body(doc, f"既往住院次数呈现最清晰的风险梯度：0 次组为 {pct(prior_stats.loc[prior_stats['prior_inpatient_capped'] == 0, 'mean'].iloc[0], 2)}，1 次组为 {pct(prior_stats.loc[prior_stats['prior_inpatient_capped'] == 1, 'mean'].iloc[0], 2)}，5 次及以上组达到 {pct(prior_stats.loc[prior_stats['prior_inpatient_capped'] == 5, 'mean'].iloc[0], 2)}。这一变量同时代表病情复杂度、慢病控制困难和医疗服务高利用，是后续模型中最稳定的风险信号。")

    doc.add_heading("4.3  入院与出院路径", level=2)
    add_figure(doc, "05_readmission_by_care_path.png", 15.0, "图 7  入院类型与出院去向的再入院率差异")
    add_body(doc, "入院类型与出院去向关联患者病情和后续照护需求。急诊入院通常比择期入院风险更高；出院至机构或转院、居家支持服务与自行离院（AMA）也代表不同的照护连续性。由于预测时点设在出院时，本文允许使用出院去向；但这些差异属于相关而非因果，不能据此断言某种出院安排导致再入院。")

    doc.add_heading("4.4  实验室检查与护理强度", level=2)
    add_figure(doc, "07_a1c_readmission.png", 12.8, "图 8  HbA1c 检测状态与再入院率")
    add_body(doc, "A1Cresult 缺失率高达 83.28%，简单删除会损失大量样本。将其拆分为“是否检测”和“是否异常”后，可以区分流程信息与检验结果。检测状态与再入院率存在差异，但这种差异可能同时反映医院流程、患者复杂度和选择性检测。")
    add_figure(doc, "06_correlation_heatmap.png", 14.0, "图 9  数值及工程特征的 Spearman 相关矩阵")
    add_body(doc, "相关矩阵显示，住院天数、药物数量、检查次数和护理强度之间存在显著共线性；既往门诊、急诊和住院次数也共同刻画医疗利用。逻辑回归需要标准化与正则化来缓解共线性，树模型则可学习非线性阈值，但可能把高度相关变量的贡献分散。")

    doc.add_heading("5  特征工程方案", level=1)
    doc.add_heading("5.1  三层对照设计", level=2)
    add_table(doc, ["方案", "输入列数", "主要内容", "研究目的"], [
        ("Raw", "41", "基础数值、原始年龄区间、三类 ID、实验室状态、23 个药物状态", "模拟直接使用原始业务字段"),
        ("Cleaned", "43", "年龄中点、语义入院/出院组、专科低频合并、payer 可得性", "检验语义清洗和缺失编码"),
        ("Engineered", "66", "Cleaned + 诊断大类、药物聚合、既往利用、强度比率、检测指示", "检验领域衍生变量的边际价值"),
    ], [2.4, 2.1, 7.4, 4.0])
    add_caption(doc, "表 4  三套特征方案及其对照逻辑")
    add_body(doc, "三套方案使用相同的患者划分、模型超参数和评价方法。这样，性能差异可以更接近地归因于特征表示，而不是样本变化或模型调参。Raw 并非“完全未经处理”：类别仍需编码、数值仍需填补；其含义是尽量保留原始字段表达。")

    doc.add_heading("5.2  领域特征构造", level=2)
    add_table(doc, ["特征组", "代表变量", "构造方法与假设"], [
        ("人口学", "age_mid", "年龄区间中点，允许树模型学习非线性阈值"),
        ("护理路径", "admission_type_group、discharge_group", "把行政 ID 映射为急诊、转院、居家支持等语义组"),
        ("诊断", "diag1_group-3_group", "ICD-9 映射为糖尿病、循环、呼吸、消化等大类"),
        ("药物", "active_med_count、dose_change_count", "统计活跃药物与 Up/Down 剂量变化次数"),
        ("既往利用", "prior_visits_total、prior_inpatient_flag", "汇总过去一年门诊、急诊和住院使用"),
        ("护理强度", "labs_per_day、care_intensity", "按住院天数标准化检查、操作与药物数量"),
        ("实验室", "A1C_tested、A1C_abnormal", "把高缺失字段拆为是否检测与是否异常"),
    ], [2.5, 5.0, 8.4])
    add_caption(doc, "表 5  领域特征工程及其解释")

    doc.add_heading("5.3  核心实现示例", level=2)
    add_code(doc, r'''
df["prior_visits_total"] = df[[
    "number_outpatient", "number_emergency", "number_inpatient"
]].sum(axis=1)

days = df["time_in_hospital"].clip(lower=1)
df["labs_per_day"] = df["num_lab_procedures"] / days
df["medications_per_day"] = df["num_medications"] / days
df["care_intensity"] = (
    df["num_lab_procedures"]
    + 2 * df["num_procedures"]
    + df["num_medications"]
) / days

df["A1C_tested"] = df["A1Cresult"].ne("Missing").astype(int)
df["A1C_abnormal"] = df["A1Cresult"].isin([">7", ">8"]).astype(int)
''')
    add_note(doc, "工程原则", "衍生变量必须在预测时点可获得、有清晰业务含义，并在相同数据划分上验证边际价值。仅因为一个变量“看起来复杂”并不代表它能提高泛化性能。")

    doc.add_heading("6  模型与实验设计", level=1)
    doc.add_heading("6.1  两类模型", level=2)
    add_body(doc, "逻辑回归是可解释线性基线。数值变量采用中位数填补与标准化，类别变量采用高频合并后的 One-Hot 编码，L2 正则化抑制高维稀疏系数。它能够检验特征工程是否把非线性或稀疏信息转化为可线性利用的信号。")
    add_body(doc, "HistGradientBoosting 是非线性提升树模型。数值变量中位数填补，类别变量使用未知值安全的序数编码；模型通过连续分裂学习阈值与交互，使用早停、叶节点最小样本数和 L2 正则控制过拟合。与随机森林相比，它在 10 万级表格数据上更高效，也更适合比较丰富的数值派生特征。")
    add_table(doc, ["项目", "逻辑回归", "HistGradientBoosting"], [
        ("主要作用", "线性可解释基线", "非线性和交互主模型"),
        ("类别处理", "One-Hot + 低频合并", "OrdinalEncoder + 未知值 -1"),
        ("数值处理", "中位数填补 + 标准化", "中位数填补，无需标准化"),
        ("复杂度控制", "C=0.7，L2 正则", "220 轮、31 叶节点、早停、L2=1"),
        ("优势", "系数方向清晰、基线稳定", "学习阈值与特征交互"),
        ("局限", "线性假设较强", "类别序数编码可能引入人为顺序"),
    ], [2.6, 6.6, 6.7])
    add_caption(doc, "表 6  两类模型的设计与角色")

    doc.add_heading("6.2  类别不平衡与评价指标", level=2)
    add_body(doc, "正类率仅 11.39%，因此 ROC-AUC 之外必须报告 PR-AUC。PR-AUC 衡量精确率—召回率曲线下的面积，其随机基准约为正类率，更能揭示少数类识别能力。概率质量用 Brier score 和校准曲线检查；阈值指标包括召回率、精确率、特异度、F1 与 F2。")
    add_body(doc, "筛查任务更重视漏诊，因此用 β=2 的 F2 分数使召回率权重高于精确率。每个模型在验证集上从 0.03 到 0.50 搜索 F2 最大阈值，之后锁定阈值并一次性评价测试集。该设计避免根据测试集结果反复调整阈值。")
    add_code(doc, r'''
def choose_f2_threshold(y_true, probability):
    thresholds = np.linspace(0.03, 0.50, 189)
    scores = [
        fbeta_score(y_true, probability >= t, beta=2, zero_division=0)
        for t in thresholds
    ]
    best = int(np.argmax(scores))
    return thresholds[best], scores[best]
''')

    doc.add_heading("6.3  稳健性与解释性分析", level=2)
    add_list(doc, [
        "用患者组自助法（200 次）估计最佳方案 ROC-AUC 与 PR-AUC 的 95% 区间，并对所选方案与 Raw 方案的 PR-AUC 差进行配对区间估计。",
        "在独立测试集子样本上计算 permutation importance，以 PR-AUC 下降衡量每个原始输入列的全局贡献。",
        "比较假阴性、假阳性与正确预测样本的年龄、既往住院、住院天数和诊断数量。",
        "按性别、种族与年龄组报告 ROC-AUC、PR-AUC、召回率和精确率，作为描述性公平性审计。",
    ])

    doc.add_heading("7  实验结果与特征工程影响", level=1)
    doc.add_heading("7.1  六组模型的排序性能", level=2)
    add_figure(doc, "09_model_comparison.png", 15.6, "图 10  三套特征方案在两类模型上的 ROC-AUC 与 PR-AUC")
    result_rows = []
    for _, row in models.sort_values(["model", "feature_set"]).iterrows():
        result_rows.append((
            row["feature_set"], row["model"].replace("HistGradientBoosting", "HistGB"),
            dec(row["roc_auc"]), dec(row["pr_auc"]), pct(row["recall"]),
            pct(row["precision"]), dec(row["f2"]), dec(row["threshold"], 3),
        ))
    add_table(doc, ["特征", "模型", "ROC-AUC", "PR-AUC", "召回率", "精确率", "F2", "阈值"],
              result_rows, [1.7, 2.3, 1.9, 1.8, 1.8, 1.8, 1.4, 1.2],
              alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 8, font_size=8.2)
    add_caption(doc, "表 7  六组患者级测试结果（阈值由验证集选择）")
    add_body(doc, "HistGradientBoosting 在三套特征上均优于逻辑回归，说明非线性阈值和交互对该任务重要。Raw + HistGB 的 PR-AUC 最高（0.210），Cleaned + HistGB 的 ROC-AUC 最高（0.664）；二者差异很小。Engineered + HistGB 的 PR-AUC 为 0.209，没有超过 Raw，说明额外诊断与强度特征没有稳定增加排序信息。")

    doc.add_heading("7.2  特征工程的收益与代价", level=2)
    add_figure(doc, "14_threshold_tradeoff.png", 13.8, "图 11  HistGradientBoosting 在三套特征下的阈值指标权衡")
    add_body(doc, "若只看 PR-AUC，会选择 Raw；若按照预先设定的验证集 F2 选择筛查方案，则 Cleaned + HistGB 最优。其测试集召回率为 81.54%，比 Raw + HistGB 的 73.62% 高 7.92 个百分点；F2 从 0.416 增至 0.418。但精确率从 15.18% 降至 14.20%，特异度从 48.32% 降至 38.13%，意味着每识别更多高风险患者，也会产生更多需人工复核的假阳性。")
    add_note(doc, "关键结论", "特征工程改变的不只是 AUC，还会改变概率分布和最佳决策阈值。Cleaned 方案的价值体现在高召回筛查，而不是显著提高排序能力。")

    doc.add_heading("7.3  ROC、PR 与阈值表现", level=2)
    add_figure(doc, "10_roc_pr_curves.png", 15.3, "图 12  HistGradientBoosting 的 ROC 与 Precision-Recall 曲线")
    add_figure(doc, "11_confusion_calibration.png", 14.8, "图 13  所选模型的混淆矩阵与概率校准")
    best = summary["best_metrics"]
    add_body(doc, f"所选 Cleaned + HistGB 的验证集阈值为 {best['threshold']:.3f}。在 {split['test_rows']:,} 条测试记录上，真阳性 1,811、假阴性 410、真阴性 6,743、假阳性 10,942；召回率 {pct(best['recall'], 2)}、精确率 {pct(best['precision'], 2)}、特异度 {pct(best['specificity'], 2)}、F2 {best['f2']:.3f}。较低阈值符合筛查用途，却不能直接作为自动干预决策。")
    add_body(doc, f"患者组自助法得到 ROC-AUC 95% 区间 [{summary['bootstrap_ci']['roc_auc'][0]:.3f}, {summary['bootstrap_ci']['roc_auc'][1]:.3f}]，PR-AUC 区间 [{summary['bootstrap_ci']['pr_auc'][0]:.3f}, {summary['bootstrap_ci']['pr_auc'][1]:.3f}]。所选方案与 Raw 方案的 PR-AUC 差区间为 [{summary['bootstrap_ci']['pr_auc_delta_selected_vs_raw'][0]:.3f}, {summary['bootstrap_ci']['pr_auc_delta_selected_vs_raw'][1]:.3f}]，跨越 0，说明排序差异没有稳定证据。")

    doc.add_heading("7.4  五折稳定性与有限容量评估", level=2)
    add_figure(doc, "15_group_cv_stability.png", 14.2, "图 14  Cleaned + HistGB 的患者分组五折交叉验证")
    cv_rows = [
        (f"第 {int(row['fold'])} 折", f"{int(row['n']):,}", pct(row["positive_rate"], 2),
         dec(row["roc_auc"], 4), dec(row["pr_auc"], 4), dec(row["brier"], 4))
        for _, row in group_cv.iterrows()
    ]
    add_table(doc, ["折数", "验证记录", "正类率", "ROC-AUC", "PR-AUC", "Brier"], cv_rows,
              [2.0, 2.6, 2.4, 2.7, 2.7, 2.7], alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 6)
    add_caption(doc, "表 8  患者互斥的五折交叉验证结果")
    ext = summary["extended_experiments"]
    add_body(doc, f"五个验证折的 ROC-AUC 均值为 {ext['cv_roc_mean']:.4f}±{ext['cv_roc_sd']:.4f}，PR-AUC 为 {ext['cv_pr_mean']:.4f}±{ext['cv_pr_sd']:.4f}。折间波动较小，表明主要结论不依赖某一次划分；其均值与独立测试集结果相近，也为模型的内部稳定性提供了另一组证据。")

    add_figure(doc, "16_capacity_curve.png", 14.2, "图 15  有限随访容量下的病例捕获率与 Lift")
    cap_rows = [
        (pct(row["capacity"], 0), f"{int(row['selected_n']):,}", f"{int(row['true_positive']):,}",
         pct(row["precision"], 2), pct(row["recall_capture"], 2), f"{row['lift']:.2f}×")
        for _, row in capacity.iterrows()
    ]
    add_table(doc, ["筛查容量", "筛查数", "命中正类", "精确率", "病例捕获率", "Lift"], cap_rows,
              [2.3, 2.5, 2.6, 2.5, 3.2, 2.4], alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 6)
    add_caption(doc, "表 9  按风险排序的 top-k 容量分析")
    add_body(doc, "当随访资源只能覆盖测试集中风险最高的 10% 记录时，可捕获 22.11% 的 30 天再入院病例，精确率 24.66%，相对随机筛查的 Lift 为 2.21。扩展到 20% 容量可捕获 38.50% 病例。这种表达比单一阈值更接近医院的人力预算决策。")

    doc.add_heading("7.5  关键特征消融与泄漏对照", level=2)
    add_figure(doc, "17_feature_ablation.png", 13.5, "图 16  移除关键特征后的 PR-AUC 变化")
    abl_rows = [
        (row["experiment"], "—" if pd.isna(row["removed"]) else row["removed"], dec(row["roc_auc"], 3),
         dec(row["pr_auc"], 3), f"{row['pr_auc_change']:+.3f}")
        for _, row in ablation.iterrows()
    ]
    add_table(doc, ["实验", "移除变量", "ROC-AUC", "PR-AUC", "PR 变化"], abl_rows,
              [3.0, 6.0, 2.6, 2.6, 2.6], font_size=8.5)
    add_caption(doc, "表 10  关键特征消融实验")
    add_body(doc, "移除“过去一年住院次数”后，PR-AUC 从 0.208 降至 0.170；同时移除既往住院、出院去向和支付信息可得性时，PR-AUC 降至 0.149。这与置换重要性结论互相印证，也说明主要信号来自患者既往利用和当次护理路径。")

    add_figure(doc, "18_split_leakage.png", 12.8, "图 17  患者级切分与随机记录切分的对照")
    leak_rows = [
        (row["split"], dec(row["roc_auc"]), dec(row["pr_auc"]), pct(row["recall"]),
         pct(row["precision"]), f"{int(row['overlap_patients']):,}")
        for _, row in leakage.iterrows()
    ]
    add_table(doc, ["切分方式", "ROC-AUC", "PR-AUC", "召回率", "精确率", "重叠患者"], leak_rows,
              [3.3, 2.3, 2.3, 2.4, 2.4, 3.0], alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 6)
    add_caption(doc, "表 11  数据切分泄漏对照")
    add_body(doc, f"若按住院记录随机切分，将有 {int(ext['row_split_overlap_patients']):,} 名患者同时出现在训练与测试集，PR-AUC 从患者互斥评估的 0.208 升至 {ext['row_split_pr_auc']:.3f}。这一对照实证了随机记录切分的乐观偏差，并说明本文的患者级互斥划分是必要的泄漏控制。")

    doc.add_heading("8  模型解释、误差与亚组审计", level=1)
    doc.add_heading("8.1  全局特征重要性", level=2)
    add_figure(doc, "12_permutation_importance.png", 13.7, "图 18  Cleaned + HistGB 的全局置换重要性")
    top_rows = []
    cn = {
        "number_inpatient": "过去一年住院次数", "discharge_group": "出院去向组",
        "payer_available": "支付信息是否记录", "insulin": "胰岛素状态",
        "number_emergency": "过去一年急诊次数", "age_mid": "年龄中点",
        "time_in_hospital": "住院天数", "number_diagnoses": "诊断数量",
        "A1Cresult": "HbA1c 结果", "diabetesMed": "是否使用糖尿病药物",
    }
    for _, row in importance.head(10).iterrows():
        top_rows.append((row["feature"], cn.get(row["feature"], row["feature"]), f"{row['importance_mean']:.4f}", f"±{row['importance_std']:.4f}"))
    add_table(doc, ["变量", "中文含义", "PR-AUC 平均下降", "重复标准差"], top_rows, [4.4, 5.1, 3.4, 3.0])
    add_caption(doc, "表 12  置换重要性前 10 位")
    add_body(doc, "既往住院次数的置换影响约为 0.061，明显高于其他变量；出院去向约为 0.020。支付信息可得性排名第三，不应被解释为支付方式的因果作用，它可能代理医院数据完整性、保险流程或患者社会经济差异。重要性接近 0 的药物变量并不意味着药物无临床价值，而是其在现有特征和数据口径下缺乏额外预测贡献。")

    doc.add_heading("8.2  错误分析", level=2)
    err_rows = []
    err_cn = {"Correct": "预测正确", "False negative": "假阴性", "False positive": "假阳性"}
    for _, row in errors.iterrows():
        err_rows.append((
            err_cn[row["error_type"]], f"{int(row['count']):,}", f"{row['mean_probability']:.3f}",
            f"{row['mean_age']:.1f}", f"{row['mean_prior_inpatient']:.2f}",
            f"{row['mean_time_in_hospital']:.2f}", f"{row['mean_diagnoses']:.2f}",
        ))
    add_table(doc, ["类型", "记录数", "平均概率", "年龄", "既往住院", "住院天数", "诊断数"],
              err_rows, [2.1, 2.0, 2.0, 1.8, 2.3, 2.6, 2.2],
              alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 7)
    add_caption(doc, "表 13  所选阈值下的错误类型特征")
    add_body(doc, "假阴性只有 410 条，符合高召回目标。它们的平均既往住院次数仅 0.002，说明模型最容易漏掉缺乏历史利用信号、却在本次出院后快速再入院的患者。假阳性平均年龄 69.2 岁、既往住院 0.89 次、住院 5.02 天、诊断 7.96 个，更像“临床上确实复杂但未在 30 天内发生结局”的人群。对筛查系统而言，这些假阳性未必完全无价值，但会增加人工复核和干预成本。")

    doc.add_heading("8.3  亚组表现与公平性", level=2)
    add_figure(doc, "13_subgroup_performance.png", 14.2, "图 19  主要种族组的召回率与精确率")
    subgroup_rows = []
    selected_groups = subgroups[
        ((subgroups["variable"] == "gender"))
        | ((subgroups["variable"] == "race_clean") & subgroups["group"].isin(["AfricanAmerican", "Caucasian", "Hispanic", "Missing"]))
    ]
    for _, row in selected_groups.iterrows():
        subgroup_rows.append((
            row["variable"], row["group"], f"{int(row['n']):,}", pct(row["positive_rate"]),
            dec(row["roc_auc"]), dec(row["pr_auc"]), pct(row["recall"]), pct(row["precision"]),
        ))
    add_table(doc, ["变量", "组别", "n", "正类率", "ROC", "PR", "召回", "精确"], subgroup_rows,
              [1.8, 3.1, 1.7, 1.8, 1.6, 1.6, 1.6, 1.7],
              alignments=[WD_ALIGN_PARAGRAPH.CENTER] * 8, font_size=8.2)
    add_caption(doc, "表 14  主要性别与种族亚组测试表现")
    add_body(doc, "女性召回率 82.70%，男性 80.21%，差异较小；Caucasian 组召回率 82.88%，AfricanAmerican 组 78.78%。Hispanic 组 ROC-AUC 较高，但样本仅 437 条，区间不确定性更大。缺失种族组召回率为 67.50%，提示数据完整性本身可能影响模型表现。亚组差异需要进一步用多次外部验证和置信区间评估，不能凭单次测试集下结论。")

    doc.add_heading("9  讨论与局限", level=1)
    doc.add_heading("9.1  特征工程的真正作用", level=2)
    add_body(doc, "本研究没有得到“特征越多性能越好”的简单结论。Raw 特征已包含大量药物状态和行政编码，梯度提升可以直接利用其中部分规律；语义清洗提高了阈值筛查召回，却没有提高 PR-AUC；完整工程特征又加入高度相关的强度比率和诊断汇总，可能增加冗余或放大序数编码的局限。特征工程的价值因此体现在可解释性、数据质量和决策目标匹配，而不只是 AUC 数值。")
    add_body(doc, "另一个重要认识是评价指标决定模型选择。若医院只能对少量患者进行高成本干预，应更重视精确率或固定资源下的 top-k 命中率，Raw 模型可能更合适；若干预成本低、漏掉高风险患者代价高，Cleaned 模型的 81.54% 召回率更有吸引力。单一“最佳模型”必须依赖具体工作流程。")

    doc.add_heading("9.2  研究局限", level=2)
    add_list(doc, [
        "数据来自 1999-2008 年美国医疗系统，支付方式、诊疗指南和住院管理已经变化；模型不能直接迁移到当前中国医院。",
        "目标变量记录的是是否再入院，不区分计划性与非计划性，也不能确认是否在数据覆盖网络之外就诊。",
        "预测时点设为出院时，因此使用出院去向与完整住院过程；若需要入院初期预测，必须重新设计特征集并重训。",
        "患者级随机划分防止个体泄漏，但不能替代医院级或时间外部验证；数据中没有明确医院 ID，无法检验跨机构泛化。",
        "类别变量在 HistGradientBoosting 中采用序数编码，虽然树模型可缓解线性顺序假设，但原生类别提升算法可能表现更好。",
        "模型是关联预测，不证明既往住院、支付信息或胰岛素使用会导致再入院；干预设计仍需要临床研究。",
        "亚组分析为描述性审计，部分少数组样本量较小，未进行正式公平约束或多重比较校正。",
    ])

    doc.add_heading("9.3  可进一步扩展的方向", level=2)
    add_list(doc, [
        "进行时间外验证或医院外验证，并检验 1999-2003 训练、2004-2008 测试的分布漂移。",
        "使用 CatBoost 等原生类别模型，比较目标编码、频率编码与当前 One-Hot/Ordinal 表示。",
        "对概率做 isotonic 或 Platt 校准，并用本地随访成本与净收益对 top-k 容量做前瞻性调参。",
        "加入社会经济、出院后随访、药物依从性和疾病严重度指标，提高可干预性与因果解释价值。",
        "建立成本敏感决策曲线，量化一次额外随访与避免一次再入院之间的净收益。",
    ])

    doc.add_heading("10  结论与建议", level=1)
    add_body(doc, "本文围绕“特征工程如何影响机器学习性能”完成了从数据审计、探索性分析、特征构造、患者级实验、阈值选择到解释与公平性审计的完整流程。99,340 条记录中，30 天再入院率为 11.39%，目标高度不平衡；既往住院次数呈现最强风险梯度。HistGradientBoosting 整体优于逻辑回归，但三套特征方案的影响并非单向。")
    add_body(doc, "按照验证集 F2 选择的 Cleaned + HistGradientBoosting 在测试集实现 ROC-AUC 0.664、PR-AUC 0.208、召回率 81.54%。它适合用作低阈值筛查工具：先识别较大的候选高风险人群，再由临床人员结合病情、社会支持与干预资源复核。若资源有限，应提高阈值或采用 top-k 排序，以减少假阳性。")
    add_body(doc, "最重要的结论不是某个模型胜出，而是研究设计本身：在医疗表格数据中，明确预测时点、保留缺失语义、把行政编码转换为业务含义、按患者划分数据并使用适合不平衡任务的评价指标，比盲目增加特征或追求单一准确率更重要。")
    add_note(doc, "最终建议", "将模型定位为出院时风险排序与随访资源分配辅助工具；上线前必须完成本地医院、近期时间段和不同人群的外部验证，并建立人工复核与持续漂移监测。", color=CORAL)

    doc.add_heading("参考文献", level=1)
    refs = [
        "[1] UCI Machine Learning Repository. Diabetes 130-US Hospitals for Years 1999-2008. DOI: 10.24432/C5230J.",
        "[2] Strack, B., et al. Impact of HbA1c Measurement on Hospital Readmission Rates: Analysis of 70,000 Clinical Database Patient Records. BioMed Research International, 2014:781670. DOI: 10.1155/2014/781670.",
        "[3] Friedman, J. H. Greedy Function Approximation: A Gradient Boosting Machine. Annals of Statistics, 29(5):1189-1232, 2001.",
        "[4] Saito, T., and Rehmsmeier, M. The Precision-Recall Plot Is More Informative than the ROC Plot When Evaluating Binary Classifiers on Imbalanced Datasets. PLOS ONE, 10(3):e0118432, 2015.",
        "[5] Pedregosa, F., et al. Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12:2825-2830, 2011.",
        "[6] Steyerberg, E. W. Clinical Prediction Models. Springer, 2nd edition, 2019.",
        "[7] Van Calster, B., et al. Calibration: The Achilles Heel of Predictive Analytics. BMC Medicine, 17:230, 2019.",
        "[8] Powers, D. M. W. Evaluation: From Precision, Recall and F-Measure to ROC, Informedness, Markedness and Correlation. Journal of Machine Learning Technologies, 2(1):37-63, 2011.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(ref)
        set_run_font(r, size=9.2)

    doc.add_page_break()
    doc.add_heading("附录 A  复现环境与运行方式", level=1)
    code_hash = hashlib.sha256(SOURCE_CODE.read_bytes()).hexdigest()
    add_table(doc, ["项目", "值"], [
        ("Python", summary["versions"]["python"]),
        ("pandas", summary["versions"]["pandas"]),
        ("NumPy", summary["versions"]["numpy"]),
        ("scikit-learn", summary["versions"]["scikit_learn"]),
        ("随机种子", "42"),
        ("代码 SHA-256", code_hash),
    ], [4.0, 11.9])
    add_caption(doc, "表 A1  本次实验的软件环境与代码校验信息")
    add_body(doc, "完整代码已单独保存为 TXT，满足课程提交要求。数据文件与代码放在同一目录或保持脚本中的 DATA_DIR 路径后，使用已配置的 torch-m4 环境运行：")
    add_code(doc, r'''
/Users/huahaowen/miniforge3/envs/torch-m4/bin/python \
    diabetes_readmission_project.py
''')
    add_list(doc, [
        "输入：diabetic_data.csv、IDS_mapping.csv。脚本不修改原始文件。",
        "输出：19 张图、12 张分析表、数据审计 JSON、划分摘要与项目结果摘要。",
        "运行流程：清洗与特征构造 → EDA → 患者级划分 → 六组模型 → 阈值优化 → 自助法与五折验证 → 容量、消融与泄漏对照 → 解释与亚组分析。",
        "代码中的绝对路径可按本机目录修改；其余随机状态固定为 42。",
    ])

    doc.add_heading("附录 B  代码文件结构", level=1)
    add_table(doc, ["模块", "核心函数", "作用"], [
        ("数据准备", "prepare_data", "缺失审计、样本排除、语义映射、特征构造"),
        ("探索分析", "make_eda_figures", "目标、缺失、年龄、利用、路径、A1c、相关性图表"),
        ("数据划分", "patient_level_split", "按 patient_nbr 分层互斥划分"),
        ("模型管道", "make_pipeline", "逻辑回归与 HistGradientBoosting 的预处理和训练"),
        ("模型评价", "choose_f2_threshold / metric_row", "验证集阈值与完整测试指标"),
        ("稳健性", "bootstrap_group_ci", "患者组自助法置信区间"),
        ("扩展验证", "extended_experiments", "五折验证、top-k 容量、特征消融与泄漏对照"),
        ("解释审计", "explain_best_model", "置换重要性、错误分析、亚组表现"),
    ], [2.4, 5.3, 8.2])
    add_caption(doc, "表 B1  完整代码的主要模块")
    add_note(doc, "提交提醒", "课程要求研究报告首页包含姓名、学号，并另交一份包含所有代码的 TXT。请打印报告与代码文件，按顺序装订，于 7 月 1 日前交至未来科创中心 601 范老师座位。", color=CORAL)

    doc.save(REPORT)
    CODE_TXT.write_text(SOURCE_CODE.read_text(encoding="utf-8"), encoding="utf-8")
    print(REPORT)
    print(CODE_TXT)


if __name__ == "__main__":
    build_report()
